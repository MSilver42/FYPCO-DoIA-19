import xlrd
import easygui
from docx import Document

from pathlib import Path


loc = Path(easygui.fileopenbox(title="Select excel file",filetypes=["*.xlsx", "*.xlsm"]))
wb = xlrd.open_workbook(loc)
sheet = wb.sheet_by_index(0)
nrows = sheet.nrows
ncols = sheet.ncols
ands = ["seconded","thirded","fourthed","fithed","sixthed","seventhed","eighthed","ninethed"]


doc = Document()  # A file name can be put in these brackets
doc.add_paragraph('thing 1', style='List')
doc.add_paragraph('thing 2', style='List Number 2')


doc.add_page_break()

#Motion list starts here
doc.add_heading("Motion List", level=0)
for i in range(1, nrows-1):
    para = doc.add_paragraph("")
    para.add_run('Motion %s' % (i)).underline = True
    para.underline = False
    para.add_run("\nWhereas:\t%s" % (sheet.cell_value(i, 0)))
    wheres = sheet.cell_value(i, 1).split(":")
    for j in wheres:
        para.add_run("\n& whereas:\t%s" % (j))
    resolutions = sheet.cell_value(i, 2).split(":")
    para.add_run("\n\nBE IT RESOLVED THAT:")
    para.add_run("\n%s" % (resolutions[0]))
    if len(resolutions)>1:
        for l in resolutions[1:]:
            para.add_run("\n\n& BE IT FURTHER RESOLVED THAT:")
            para.add_run("\n%s" % l)
    para.add_run("\n\n\tMoved by: %s" % (sheet.cell_value(i, 3)))
    movers = sheet.cell_value(i, 4).split(":")
    for k in range(len(movers)):
        para.add_run("\n\t%s by: %s" % (ands[k], movers[k]))
#motion list ends here

doc.save('agenda test.docx')